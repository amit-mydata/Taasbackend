import uuid
import asyncio
from bson import ObjectId
from app.utils.llm import generate_quiz_with_gemini, generate_interview_questions, generate_interview_text_questions_questions

from docx import Document
def extract_text_and_tables(file_path: str) -> str:
    """Extract paragraphs and tables from a DOCX file."""
    doc = Document(file_path)
    text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text.append("\t".join(row_data))

    return "\n".join(text)

async def process_quiz_questions(candidate_id: str, job_description: str, extracted_text: str):
    """
    Runs in background: generate quiz questions and save to DB (or log).
    """
    try:
        print("Generating Quiz Questions...")
        quiz_response, interview_questions, text_questions = await asyncio.gather(
            generate_quiz_with_gemini(job_description, extracted_text),
            generate_interview_questions(job_description, extracted_text),
            generate_interview_text_questions_questions(job_description, extracted_text)
        )

        print("Generated Quiz Response:", quiz_response)  
        print("Generated Interview Questions:", interview_questions)  
        print("Generated Interview Text Questions:", text_questions)

        quiz_list = []
        # Collect all quiz items
        if candidate_id and quiz_response and "quiz" in quiz_response:
            quiz_list = []
            for quiz_item in quiz_response["quiz"]:
                if isinstance(quiz_item, dict):
                    quiz_data = {
                        "quiz_id": str(uuid.uuid4()),
                        "question": quiz_item.get("question", quiz_item),
                        "options": quiz_item.get("options", []),
                        "correct_answer": quiz_item.get("correct_answer", None),
                        "type": "mcqs_questions"
                    }
                quiz_list.append(quiz_data)

        
        # Add interview questions to quiz_list (with answer field)
        if candidate_id and interview_questions and "questions" in interview_questions:
            for qa in interview_questions["questions"]:
                quiz_data = {
                    "quiz_id": str(uuid.uuid4()),
                    "question": qa.get("question", ""),
                    "correct_answer": qa.get("answer", None), 
                    "type": "coding_questions"              
                }
                quiz_list.append(quiz_data)

        if candidate_id and text_questions and "questions" in text_questions:
            for qa in text_questions["questions"]:
                quiz_data = {
                    "quiz_id": str(uuid.uuid4()),
                    "question": qa.get("question", ""),
                    "correct_answer": qa.get("answer", None),
                    "type": "text_questions"
                }
                quiz_list.append(quiz_data)

        quiz_list = [{'quiz_id': '3f5059a6-c63b-47b9-a29c-1bd15f7852f2', 'question': "You are tasked with creating an IT Hardware Asset budget forecast, a task you've mentioned in your resume. Which data-driven approach would you prioritize for the initial analysis?", 'options': ["Primarily using the previous year's budget as a fixed baseline for the new forecast.", 'Analyzing historical procurement data, current asset refresh cycles, and projected business growth.', 'Surveying department heads for their hardware wish-list items without historical context.', 'Focusing solely on the cheapest available hardware options currently on the market.'], 'correct_answer': 'Analyzing historical procurement data, current asset refresh cycles, and projected business growth.', 'type': 'mcqs_questions'}, {'quiz_id': 'f47e2ef2-8593-47ac-beed-0fa865b3334e', 'question': 'When creating a monthly KPI report for asset accuracy, as you did at Infosys, what is the most critical data comparison you would include to measure the effectiveness of your CMDB?', 'options': ['The number of new assets purchased versus the number of assets disposed of during the month.', 'The total financial value of all assets currently in the inventory.', 'The number of assets recorded in the CMDB versus the number of assets found by discovery tools like SCCM.', 'The total count of support tickets related to hardware failures.'], 'correct_answer': 'The number of assets recorded in the CMDB versus the number of assets found by discovery tools like SCCM.', 'type': 'mcqs_questions'}, {'quiz_id': '2738db15-881b-4007-b378-30e9ee67a997', 'question': "You've built dashboards in both ServiceNow and Excel. If you were asked to build a new, real-time ITAM dashboard for executive leadership, what would be your primary goal for the main view?", 'options': ['To display every single data field available in the raw asset database for maximum transparency.', 'To provide high-level, actionable insights like compliance status, asset lifecycle stages, and stock levels at a glance.', 'To create a visually complex chart that showcases advanced data visualization techniques.', 'To provide a detailed, scrolling list of all individual asset procurement requests and their current statuses.'], 'correct_answer': 'To provide high-level, actionable insights like compliance status, asset lifecycle stages, and stock levels at a glance.', 'type': 'mcqs_questions'}, {'quiz_id': '222e08fd-35e6-4ad4-94d3-aa02d1aae392', 'question': 'To ensure CMDB data quality and governance, what would be your first step when you identify discrepancies between your CMDB and data from an asset discovery tool?', 'options': ["Immediately overwrite the CMDB data with the data from the discovery tool to ensure it's up to date.", 'Assume the discovery tool is incorrect and make no changes to the CMDB.', 'Initiate a formal reconciliation process to investigate the source of the discrepancy and apply corrections based on defined governance rules.', 'Delete the conflicting Configuration Item (CI) from the CMDB to resolve the conflict.'], 'correct_answer': 'Initiate a formal reconciliation process to investigate the source of the discrepancy and apply corrections based on defined governance rules.', 'type': 'mcqs_questions'}, {'quiz_id': '314e5542-339b-49c7-afa9-305b624896ca', 'question': 'In your role managing Software Asset Management (SAM), how would you use data analysis to identify significant cost-saving opportunities for a major software publisher like Microsoft?', 'options': ['By simply counting the total number of licenses owned versus the number of employees in the company.', 'By comparing license entitlements against actual software usage data to identify underutilized licenses that can be harvested or retired.', 'By cross-referencing employee hire dates with the dates their software licenses were assigned.', 'By focusing only on the compliance status (compliant vs. non-compliant) without analyzing software usage patterns.'], 'correct_answer': 'By comparing license entitlements against actual software usage data to identify underutilized licenses that can be harvested or retired.', 'type': 'mcqs_questions'}, {'quiz_id': '4b794880-ac87-4c26-a7f8-48480cb2bf67', 'question': 'Given your experience in vendor management, how would you approach using data to objectively measure and compare vendor performance for IT asset procurement?', 'options': ['By tracking only the total annual amount spent with each individual vendor.', "By relying on the personal relationship you have with the vendor's account manager.", 'By analyzing vendor data on delivery lead times, hardware failure rates, and adherence to SLAs to create a performance scorecard.', 'By consistently selecting the vendor that offers the lowest absolute price for every order, regardless of other factors.'], 'correct_answer': 'By analyzing vendor data on delivery lead times, hardware failure rates, and adherence to SLAs to create a performance scorecard.', 'type': 'mcqs_questions'}, {'quiz_id': 'a10b2361-67be-447f-b919-cc3428d5dec1', 'question': 'With your ITIL V4 certification and ITSM experience, how would you use data from incident and change management tickets to proactively improve the IT Asset Management lifecycle?', 'options': ['By analyzing ticket data to identify hardware models with high failure rates, which could inform future procurement decisions.', 'By counting the total number of tickets logged each month to justify hiring more IT support staff.', 'By treating ITSM data as separate and not directly relevant to the core functions of asset management.', 'By tracking only the average resolution times for asset-related tickets without analyzing the root causes.'], 'correct_answer': 'By analyzing ticket data to identify hardware models with high failure rates, which could inform future procurement decisions.', 'type': 'mcqs_questions'}, {'quiz_id': '12e8117c-e521-444f-a263-da725d458c31', 'question': 'Your resume mentions reconciling asset database gaps. What is your understanding of the primary goal of data reconciliation in the context of IT asset management?', 'options': ['To ensure that two databases have the exact same number of records, even if the data within them differs.', 'To create a brand new database that merges all the data from the two original sources.', 'To identify and correct inconsistencies between data sources to create a single, accurate source of truth for all asset information.', 'To automatically select the database with the most recent entries as the new primary source of information.'], 'correct_answer': 'To identify and correct inconsistencies between data sources to create a single, accurate source of truth for all asset information.', 'type': 'mcqs_questions'}, {'quiz_id': '0544542d-2de5-489b-9a12-ff22d3a994a8', 'question': 'You list Power BI, Advanced Excel, and Cognos as BI tools you know. If you needed to present a detailed software license optimization analysis to management, which tool would you choose to create an interactive report with drill-down capabilities?', 'options': ['Advanced Excel, because its pivot tables are sufficient for any type of data analysis.', 'Power BI, for its strong capabilities in creating interactive dashboards and sharing accessible, visual reports.', 'A simple text-based report in an email to clearly summarize the key findings without visuals.', 'The built-in reporting module of an ITAM tool like SNOW, even if it lacks the flexibility of a dedicated BI tool.'], 'correct_answer': 'Power BI, for its strong capabilities in creating interactive dashboards and sharing accessible, visual reports.', 'type': 'mcqs_questions'}, {'quiz_id': '5d8de4e4-37cf-4f5e-9330-bfcf3518d73f', 'question': 'Based on your expertise in CI Lifecycle Management, what would you consider the most critical data-related risk that needs to be managed?', 'options': ['The risk of having too many different types of CIs defined within the CMDB.', 'The risk of failing to properly decommission CIs for retired assets, leading to inaccurate data for impact analysis and potential security gaps.', 'The risk of changing the name or identifier of a CI during its operational lifecycle.', 'The risk of not assigning every single CI to a specific business service or cost center.'], 'correct_answer': 'The risk of failing to properly decommission CIs for retired assets, leading to inaccurate data for impact analysis and potential security gaps.', 'type': 'mcqs_questions'}, {'quiz_id': '9fde8189-2d88-43bd-9cbd-0a537add742f', 'question': "In your experience with CMDB administration in ServiceNow, you've managed complex Configuration Item (CI) relationships. If you were tasked to design a data model for a new CMDB, specifically optimized for analytical and predictive reporting (e.g., forecasting hardware failures), what would you do differently compared to a standard operational CMDB design?", 'correct_answer': "For an analytics-focused CMDB, I would move beyond a purely transactional, normalized model and incorporate data warehousing principles. I'd design a star schema where a central fact table tracks events like incidents, changes, or performance metrics. This fact table would be linked to dimension tables for CIs, locations, business services, and time. Crucially, I would implement Slowly Changing Dimensions (SCDs) to track historical changes to CIs, which is vital for building accurate predictive models. This denormalized approach prioritizes fast query performance for aggregation, which is often a bottleneck in standard operational databases when used for complex analytics.", 'type': 'coding_questions'}, {'quiz_id': 'b5578a8c-e8d9-46a5-96dc-ace6241b3982', 'question': "Your resume highlights creating budget and forecast reports. Imagine a stakeholder asks you to build a model to predict the quarterly IT hardware procurement budget for the next fiscal year. Walk me through your problem-solving approach, including the data sources you'd use and the analytical techniques you'd apply.", 'correct_answer': "My approach would be a multi-step process focused on creating a robust forecast. First, I would gather historical data from multiple sources: procurement records from SAP Ariba for past purchases, asset lifecycle data from ServiceNow to understand refresh cycles, and HR system data for employee growth projections. Next, I would perform exploratory data analysis to identify trends, seasonality, and the average lifespan of different asset classes. I would then likely use time-series forecasting methods, such as ARIMA or exponential smoothing in Excel or Power BI, to create a baseline forecast. Finally, I'd enrich this baseline with qualitative inputs from business unit leaders about upcoming projects to produce a comprehensive and defensible budget plan.", 'type': 'coding_questions'}, {'quiz_id': 'bcbffd58-0d37-49bf-b02d-fe0ed2260fdf', 'question': "You mentioned creating real-time dashboards in ServiceNow and Power BI. Suppose a critical 'Software License Compliance' dashboard in Power BI has become extremely slow, taking several minutes to load. What is your systematic process for diagnosing and optimizing its performance?", 'correct_answer': "My first step would be to use Power BI's built-in Performance Analyzer to identify the slowest visuals and DAX queries. Often, the bottleneck is either an overly complex data model or inefficient DAX calculations. I would then review the data model for convoluted relationships, excessive cardinality, or bi-directional filters that could be simplified. For query optimization, I would push as much data transformation as possible back to the source database via Power Query, pre-aggregate data where feasible, and rewrite complex DAX measures to use more efficient functions like variables and iterators. If performance is still an issue, I would consider implementing an aggregated summary table to handle the most resource-intensive visuals.", 'type': 'coding_questions'}, {'quiz_id': 'f4618091-9aea-437e-a98f-5e89cec7700d', 'question': 'Drawing on your experience with diverse asset discovery tools like SCCM and BigFix, design a system for real-time monitoring and alerting of unauthorized software installations across a large enterprise. Describe the key components, data flow, and how you would handle data reconciliation.', 'correct_answer': "I would design a system with four key components: Data Ingestion, a Centralized Asset Database, a Rules Engine, and an Alerting/Reporting Layer. The ingestion component would use APIs or scheduled jobs to pull installation data from SCCM and other discovery tools in near real-time. This data would feed into a centralized database, likely built in SQL, which also contains our 'golden record' of approved software entitlements from a tool like Flexera or SNOW. The rules engine would continuously compare the discovered installations against the approved list, flagging any discrepancies. Finally, upon detecting unauthorized software, the alerting system would automatically generate a ServiceNow ticket assigned to the relevant support team and update a compliance dashboard for management visibility.", 'type': 'coding_questions'}, {'quiz_id': '25a274ea-d9fb-4b1b-ba7d-188a8557ebeb', 'question': "You have extensive experience creating KPI reports and performing data reconciliation. Can you describe the governance process and best practices you follow to ensure the accuracy and trustworthiness of a critical report, such as the monthly 'Asset Accuracy' KPI you mentioned creating for FAB Bank?", 'correct_answer': "To ensure the report's integrity, I follow a rigorous data governance process from end-to-end. It begins with clearly defining the 'Asset Accuracy' metric with stakeholders to avoid ambiguity. Next, I establish clear data lineage, documenting the exact data sources—like SCCM for discovered assets and ServiceNow as the system of record—and all transformation logic. Before publishing, I perform automated reconciliation checks against source systems to flag discrepancies and conduct a peer review of my logic and calculations. Finally, the report includes a data quality scorecard, providing transparency into any known data gaps, which builds trust with management and highlights areas for underlying process improvement.", 'type': 'coding_questions'}, {'quiz_id': '2fde4ded-8e36-47b3-bd4a-ea1ab284c432', 'question': "Your resume mentions creating a 'Hardware Asset Finance Portfolio Dashboard' using Advanced Excel at HCL. Could you describe the business problem this dashboard solved, the data sources you integrated, and how your analysis led to a measurable outcome?", 'correct_answer': 'At HCL, my client British Petroleum lacked a unified view of their hardware asset financial data, which was siloed in procurement systems and asset databases. My task was to create a centralized dashboard for tracking total cost of ownership and budget variance. I used Power Query in Excel to ingest and merge data from ServiceNow AMDB and SAP Ariba, then built pivot tables and charts to visualize metrics like budget vs. actuals and asset depreciation. This provided a real-time financial overview that improved budget forecasting accuracy by about 10% and helped identify cost-saving opportunities in our asset refresh cycle.', 'type': 'text_questions'}, {'quiz_id': '055066db-d956-4ac2-99ed-e8927581118c', 'question': "You highlight 'Data Quality & Governance' as a key skill, specifically mentioning the creation of gap analysis reports at IBM. Can you walk me through a significant data discrepancy you uncovered between your CMDB and a discovery tool like SCCM, and the data analysis steps you took to resolve it?", 'correct_answer': "While at IBM for the State Street account, we found a major discrepancy where our SCCM discovery tool was not reporting thousands of assets that were listed in our ServiceNow CMDB. I exported data from both systems and used Excel's VLOOKUP and pivot tables to perform a gap analysis, which isolated the issue to a specific misconfigured network subnet. After coordinating with the network team to fix the agent communication, I developed a reconciliation report to automatically flag such discrepancies in the future. This process resolved over 2,000 CI mismatches and improved our CMDB data accuracy KPI by 15%.", 'type': 'text_questions'}, {'quiz_id': '63685fcb-9735-4c66-8bce-45f5e9bfde94', 'question': 'At Tech Mahindra, you created real-time ITAM dashboards in ServiceNow. From a data science perspective, how did you approach defining the most critical KPIs for your client, Rayus Radiology, and what data visualization techniques did you use to make these dashboards impactful for management?', 'correct_answer': "For Rayus Radiology, I first met with stakeholders to understand their primary goals, which were cost control and compliance. We collaboratively defined key KPIs such as 'Asset Stock Level vs. Demand Forecast,' 'License Compliance Rate,' and 'Asset Age Distribution.' I then used ServiceNow's reporting module to build the dashboard, utilizing trend lines to show performance over time, bar charts to compare compliance across different software publishers, and pie charts to visualize the hardware age distribution. This gave leadership a clear, at-a-glance view to make informed decisions on procurement and asset lifecycle management.", 'type': 'text_questions'}, {'quiz_id': '8dbf83bb-c155-4569-91f0-bb4514abda0a', 'question': 'Imagine you are given a dataset of all IT asset procurement requests from the past two years. Based on your experience with budget planning and forecasting at Infosys, what data analysis approach would you take to identify cost-saving opportunities and predict future hardware needs?', 'correct_answer': 'My first step would be to perform exploratory data analysis using a tool like Power BI to clean the data and identify key variables like department, asset type, and request frequency. I would then analyze purchasing patterns to identify opportunities for bulk purchases or standardization of models to leverage volume discounts. For forecasting, I would apply time-series analysis to model historical demand trends and seasonality, allowing me to more accurately predict future hardware needs by department and quarter, thus optimizing our budget and inventory levels.', 'type': 'text_questions'}, {'quiz_id': 'fb8bcb58-d72c-44fc-a4bd-acd3c4fe4570', 'question': 'You automated the vendor shipping report process at HCL. Can you explain the data management principles you applied to successfully integrate varied data formats into the ServiceNow AMDB, and what was the quantifiable impact of this automation?', 'correct_answer': "The main challenge was that each vendor sent data in a slightly different CSV format. I first established a standardized data schema by identifying the core data points common across all reports, such as asset tag, model, and shipping date. I then used ServiceNow's transform maps to map each vendor's unique column headers to our standard schema and implemented data validation rules to flag any anomalies. The result was a fully automated workflow that saved approximately 10 hours of manual data entry per week and reduced data entry errors by over 95%, ensuring our asset tracking data was consistently accurate and timely.", 'type': 'text_questions'}] 

        # Store once in DB (import AnalyzerService lazily so Motor binds to the active event loop)
        from app.services.analyzer import AnalyzerService
        analyzer_service = AnalyzerService()
        # print("############################################3",candidate_id)
        # print("ssssssssssssssssssssssssssssssssssss",quiz_list)
        await analyzer_service.store_quiz_questions(candidate_id, quiz_list)

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error in background quiz generation: {str(e)}")
        raise


async def calculate_overall_score(resume, communication, technical):

    weight_resume=40
    # weight_comm=20
    weight_tech=60

    if resume is None:
        resume = 0
    if communication is None:
        communication = 0
    if technical is None:
        technical = 0
    
    # total_weight = weight_resume + weight_comm + weight_tech
    total_weight = weight_resume + weight_tech
    score = (
        (resume * weight_resume) +
        # (communication * weight_comm) +
        (technical * weight_tech)
    ) / total_weight
    score = round(score, 2)

    # Category mapping for fit status
    if score >= 85:
        status = "Strong Fit"
    elif score >= 70:
        status = "Potential Fit"
    else:
        status = "Not a Fit"
    return score, status


def convert_objectids(obj):
    """Recursively convert ObjectId instances to strings."""
    if isinstance(obj, list):
        return [convert_objectids(item) for item in obj]
    elif isinstance(obj, dict):
        return {key: convert_objectids(value) for key, value in obj.items()}
    elif isinstance(obj, ObjectId):
        return str(obj)
    else:
        return obj
